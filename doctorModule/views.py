import datetime
import json
import random
import string

import simplejson as simplejson
from django.contrib import auth
from django.contrib.auth.models import User
from django.core.paginator import Paginator
from django.db.models import Q
from django.forms import model_to_dict
from django.http import JsonResponse
# Create your views here.
from django.utils import timezone

# from mailmerge import MailMerge
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from doctorModule.forms import RegistrationForm, LoginForm
from doctorModule.models import Doctor, Diagnosis
from doctorModule.models import Patient


def get_all_doctor(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        nowPage = reqform['page']
        pageSize = reqform['pageSize'] if reqform['pageSize'] is not None else 8
        doc_list = Doctor.objects.filter(is_admin=0)
        p = Paginator(doc_list, pageSize)
        doc_page = p.page(nowPage)
        response['msg'] = 'success'
        response['error_num'] = 0
        response['data'] = objects_to_json(doc_page)
        response['total'] = p.count

    return JsonResponse(response)


def get_doctor_by_name(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        doc_name = reqform['docName']
        nowPage = reqform['page']
        pageSize = reqform['pageSize'] if reqform['pageSize'] is not None else 8
        docs = Doctor.objects.filter(doc_name=doc_name)
        p = Paginator(docs, pageSize)
        response['data'] = objects_to_json(p.page(nowPage))
        response['total'] = p.count
        response['msg'] = 'success'
        response['error_num'] = 0

    return JsonResponse(response)


def add_doctor_by_admin(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)

        username = reqform['doc_name']
        password = reqform['doc_tel'][-6:]
        email = reqform['doc_email']
        doc_name = reqform['doc_name']
        doc_birth = reqform['doc_birth']
        doc_sex = reqform['doc_sex']
        doc_depart = reqform['doc_depart']
        doc_hospital = 'xxx hospital'
        doc_tel = reqform['doc_tel']
        doc_description = reqform['doc_description']
        try:
            user = User.objects.create_user(username=username, password=password, email=email)
            doctor = Doctor(user=user, doc_name=doc_name, doc_birth=doc_birth, doc_sex=doc_sex, doc_depart=doc_depart,
                        doc_hospital=doc_hospital, doc_tel=doc_tel, doc_description=doc_description,
                        inform_time=timezone.now().date(), is_admin=False)
            doctor.save()
        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
            return JsonResponse(response)
        response['msg'] = 'success'
        response['error_num'] = 0

    return JsonResponse(response)


def delete_doctor_by_admin(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        doc_id = reqform['id']
        try:
            doctor=Doctor.objects.filter(id=doc_id)
            user_id =doctor.values('user_id').first()['user_id']
            User.objects.filter(id=user_id).delete()
            doctor.delete()

        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
            return JsonResponse(response)
        response['msg'] = 'success'
        response['error_num'] = 0

    return JsonResponse(response)

def reset_doctor_password(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        doc_id = reqform['id']
        try:
            doctor = Doctor.objects.filter(id=doc_id)
            user_id = doctor.values('user_id').first()['user_id']
            user = User.objects.get(id=user_id)
            user.set_password('123456')
            user.save()
        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
            return JsonResponse(response)
        response['msg'] = 'success'
        response['error_num'] = 0

    return JsonResponse(response)


def register(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        form = RegistrationForm(reqform)
        try:
            form.is_valid()
        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
            return JsonResponse(response)
        if form.is_valid():
            username = form.cleaned_data['username']
            email = form.cleaned_data['email']
            password = form.cleaned_data['password2']

            # 使用内置User自带create_user方法创建用户，不需要使用save()
            user = User.objects.create_user(username=username, password=password, email=email)

            # 如果直接使用objects.create()方法后不需要使用save()
            doctor = Doctor(user=user, doc_name=username)
            doctor.save()

            response['msg'] = 'success'
            response['error_num'] = 0
            return JsonResponse(response)

        else:
            response['msg'] = "form is not valid"
            response['error_num'] = 1

    else:
        response['msg'] = 'method is not POST'
        response['error_num'] = 1

    return JsonResponse(response)


def login(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        form = LoginForm(reqform)  # 需要使用 x-www-form-urlencoded 的request
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']

            user = auth.authenticate(username=username, password=password)

            if user is not None and user.is_active:
                auth.login(request, user)
                response['msg'] = 'success'
                response['error_num'] = 0
                data = Doctor.objects.get(user__username=username)
                user_info = User.objects.get(username=username)
                response['data'] = object_to_json(data)
                response['data']['user'] = object_to_json(user_info)

                return JsonResponse(response, safe=False)

            else:
                # 登陆失败
                response['msg'] = 'login fail'
                response['error_num'] = 1
                return JsonResponse(response)
        else:
            response['msg'] = 'valid fail'
            response['error_num'] = 1
        return JsonResponse(response)


def add_patient(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        patientName = reqform['patient_name']
        patientSex = reqform['patient_sex']
        patientBirth = reqform['patient_birth']
        patientDocId = int(reqform['patient_doc'])
        patientAddress = reqform['patient_address']
        patientPastHistory = reqform['patient_past_history']
        patientTel = reqform['patient_tel']
        patientRemark = reqform['patient_remark']
        patient_add_time = reqform['patient_add_time']
        try:
            newPatient = Patient(patient_name=patientName, patient_sex=patientSex, patient_birth=patientBirth,
                                 patient_doc=patientDocId, patient_address=patientAddress,
                                 patient_past_history=patientPastHistory, patient_tel=patientTel,
                                 patient_remark=patientRemark, patient_add_time=patient_add_time)
            newPatient.save()

            tempDiag = Diagnosis(doc_id=patientDocId, patient_id=newPatient.id, DiagTime=patient_add_time)
            tempDiag.save()

            response['msg'] = 'success'
            response['error_num'] = 0
        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
        return JsonResponse(response)


def query_patient_by_docId(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        docId = reqform['doctorId']
        nowPage = reqform['page']
        pageSize = reqform['pageSize'] if reqform['pageSize'] is not None else 8
        patientList = Diagnosis.objects.filter(doc_id=docId).order_by("-DiagTime")
        p = Paginator(patientList, pageSize)
        patient_page = p.page(nowPage)
        temp = []
        for patient in patient_page:
            patient_info = Patient.objects.get(id=patient.patient_id)
            patient_json = object_to_json(patient)
            patient_json['patient_info'] = object_to_json(patient_info)
            temp.append(patient_json)

        response['data'] = list_to_json(temp)
        response['total'] = p.count
        response['msg'] = 'success'
        response['error_num'] = 0
    return JsonResponse(response)


def query_diagnosis_by_patientId(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        patientId = reqform['patient_id']
        doctorId = reqform['doctor_id']
        data = objects_to_json(Diagnosis.objects.filter(Q(patient_id=patientId) & Q(doc_id=doctorId)))
        response['data'] = data
        response['msg'] = 'success'
        response['error_num'] = 0
    return JsonResponse(response)


def update_patient_info(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        form = {
            'patient_name': reqform['name'],
            'patient_sex': reqform['sex'],
            'patient_birth': reqform['birth'],
            'patient_address': reqform['address'],
            'patient_tel': reqform['tel'],
            'patient_past_history': reqform['pastHistory'],
            'patient_remark': reqform['remark']
        }
        patient_id = reqform['id']
        try:
            Patient.objects.filter(id=patient_id).update(**form)
        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
            return JsonResponse(response)
        response['data'] = objects_to_json(Patient.objects.filter(id=patient_id))
        response['msg'] = 'success'
        response['error_num'] = 0
    return JsonResponse(response)


def query_patient_by_patientName(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        patientName = reqform['patientName']
        doctorId = reqform['doctorId']
        nowPage = reqform['page']
        pageSize = reqform['pageSize'] if reqform['pageSize'] is not None else 8
        patient_this_doctor = Patient.objects.filter(patient_name__contains=patientName, patient_doc=doctorId)
        patient_that_doctor = Patient.objects.exclude(patient_doc=doctorId).filter(patient_name=patientName)
        if patient_this_doctor.exists():
            p = Paginator(patient_this_doctor, pageSize)
            response['data'] = objects_to_json(p.page(nowPage))
            response['total'] = p.count
            response['msg'] = 'success'
            response['error_num'] = 0
        else:
            p = Paginator(patient_that_doctor, pageSize)
            response['data'] = objects_to_json(p.page(nowPage))
            response['total'] = p.count
            response['msg'] = 'error doctor'
            response['error_num'] = 1

    return JsonResponse(response)


def check_doctor_info(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        doctorId = reqform['doctorId']
        doctor = objects_to_json(Doctor.objects.filter(id=doctorId))[0]
        if doctor['inform_time'] + datetime.timedelta(days=15) < timezone.now().date():
            for key in doctor.keys():
                if doctor[key] is None:
                    response['data'] = False
                    response['msg'] = key + ' is None'
                    response['error_num'] = 0
                    return JsonResponse(response)
        else:
            response['data'] = True
            response['msg'] = 'no need to disturb'
            response['error_num'] = 0
            return JsonResponse(response)
        response['data'] = True
        response['msg'] = 'info is complete'
        response['error_num'] = 0

    return JsonResponse(response)


def reset_inform_date(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        doctorId = reqform['doctorId']
        Doctor.objects.filter(id=doctorId).update(inform_time=timezone.now().date())
        response['msg'] = 'success'
        response['error_num'] = 0

    return JsonResponse(response)


def update_doctor_info(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)

        form = {
            'doc_name': reqform['docName'],
            'doc_sex': reqform['docSex'],
            'doc_birth': reqform['docBirth'],
            'doc_tel': reqform['docTel'],
            'doc_hospital': reqform['docHospital'] if 'docHospital' in reqform else "xxx hospital",
            'doc_depart': reqform['docDepart'],
            'doc_description': reqform['docDescription'] if 'docDescription' in reqform else ""
        }
        doctorId = reqform['id']
        try:
            Doctor.objects.filter(id=doctorId).update(**form)
        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
            return JsonResponse(response)
        response['data'] = objects_to_json(Doctor.objects.filter(id=doctorId))
        response['msg'] = 'success'
        response['error_num'] = 0

    return JsonResponse(response)


def query_doctor_by_id(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        doctorId = reqform['doctorId']
        docInfo = objects_to_json(Doctor.objects.filter(id=doctorId))[0]
        response['data'] = docInfo
        response['msg'] = 'success'
        response['error_num'] = 0

    return JsonResponse(response)


def add_diag_info(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        docId = reqform['doc_id']
        patientId = reqform['patient_id']
        ctName = reqform['ct_name']
        aiResult = reqform['ai_result']
        diagResult = reqform['diag_content']
        diagTime = reqform['diag_time']
        patientSymptm = reqform['symptom']
        diagRemark = reqform['diag_remark']
        img_list_int = []
        for str_num in diagRemark:
            num = int(str_num)
            img_list_int.append(num)
        try:
            diag_item = Diagnosis(doc_id=docId, patient_id=patientId, CTURL=ctName, AIResult=aiResult,
                                  DiagResult=diagResult, DiagTime=diagTime, patient_symptom=patientSymptm,
                                  DiagRemark=img_list_int)
            diag_item.save()
            response['msg'] = 'success'
            response['error_num'] = 0
            return JsonResponse(response)
        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
            return JsonResponse(response)

    return JsonResponse(response)


def update_diag_info(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        diagId = reqform['diag_id']
        docId = reqform['doc_id']
        patientId = reqform['patient_id']
        ctName = reqform['ct_name']
        aiResult = reqform['ai_result']
        diagResult = reqform['diag_content']
        diagTime = reqform['diag_time']
        patientSymptm = reqform['symptom']
        diagRemark = reqform['diag_remark']
        img_list_int = []
        for str_num in diagRemark:
            num = int(str_num)
            img_list_int.append(num)

        try:
            Diagnosis.objects.filter(id=diagId).update(doc_id=docId, patient_id=patientId, CTURL=ctName,
                                                       AIResult=aiResult, DiagResult=diagResult,
                                                       DiagTime=diagTime, patient_symptom=patientSymptm,
                                                       DiagRemark=img_list_int)

            response['msg'] = 'success'
            response['error_num'] = 0
            return JsonResponse(response)
        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
            return JsonResponse(response)

    return JsonResponse(response)


def get_diag_info(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        diag_id = reqform['id']
        try:
            diag_info = Diagnosis.objects.filter(id=diag_id)
            response['msg'] = 'success'
            response['error_num'] = 0
            response['data'] = objects_to_json(diag_info)
            return JsonResponse(response)
        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
            return JsonResponse(response)
    return JsonResponse(response)


def del_diag_info(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        diag_id = reqform['id']
        try:
            Diagnosis.objects.filter(id=diag_id).delete()
            response['msg'] = 'success'
            response['error_num'] = 0
            return JsonResponse(response)
        except Exception as e:
            response['msg'] = str(e)
            response['error_num'] = 1
            return JsonResponse(response)

    return JsonResponse(response)


def get_doc_diag_list(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        docId = reqform['doc_id']
        hasDiag_list = Diagnosis.objects.filter(doc_id=docId).exclude(CTURL__isnull=True).exclude(CTURL="")
        response['msg'] = 'success'
        response['error_num'] = 0
        response['data'] = objects_to_json(hasDiag_list)
    return JsonResponse(response)


def generateReport(request):
    response = {}
    if request.method == 'POST':
        reqform = simplejson.loads(request.body)
        patientName = reqform['patient_name']
        patientSex = '男' if reqform['patient_sex'] == 1 else '女'
        patientAge = reqform['patient_age']
        dTime = reqform['diagnosis_time']
        patientTel = reqform['patient_tel']
        patientPH = reqform['patient_pasthistory']
        doctorName = reqform['doctor_name']
        symptom = reqform['symptom']
        dContent = reqform['diagnosis_content']
        exportTime = reqform['export_time']
        img_index_list = reqform['img_list']  # class<list>
        ct_name = reqform['ct_name']

        img_path_list = []
        local_path = "D:/workspace/python/LCCAD/detector/"
        for index in img_index_list:
            all_img_path = local_path + "patient_imgs/" + ct_name + "/"
            target_img_path = all_img_path + "img" + str(index).zfill(3) + ".jpg"
            img_path_list.append(target_img_path)

        target_path = "D:/workspace/python/LCCAD/doctorModule/report/"
        ran_str = ''.join(random.sample(string.ascii_letters + string.digits, 5))
        report_name = "report" + patientName + '-' + ran_str + ".docx"
        target_report = target_path + report_name

        design_report(patient_name=patientName, patient_sex=patientSex, patient_age=patientAge, diagnosis_time=dTime,
                      patient_tel=patientTel, patient_past_history=patientPH, doc_name=doctorName, symptom=symptom,
                      img_path_list=img_path_list, diagnosis_content=dContent, save_name=target_report)

    return JsonResponse(response)


def design_report(patient_name, patient_sex, patient_age, diagnosis_time, patient_tel, patient_past_history, doc_name,
                  symptom, img_path_list, diagnosis_content, save_name):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    head = document.add_heading('', 0)
    head.add_run('xxx 医院肺科诊断结果').font.color.rgb = RGBColor(0xff, 0x00, 0x00)
    head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(7, 6, style='Table Grid')
    table.autofit = False
    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(1, 2).merge(table.cell(1, 3))
    table.cell(2, 0).merge(table.cell(2, 1))
    table.cell(2, 2).merge(table.cell(2, 5))
    table.cell(3, 0).merge(table.cell(3, 1))
    table.cell(3, 2).merge(table.cell(3, 5))
    table.cell(4, 0).merge(table.cell(4, 1))
    table.cell(4, 2).merge(table.cell(4, 5))
    table.cell(5, 0).merge(table.cell(5, 5))
    table.cell(6, 0).merge(table.cell(6, 1))
    table.cell(6, 2).merge(table.cell(6, 5))
    cells_line1 = table.rows[0].cells
    cells_line2 = table.rows[1].cells
    cells_line3 = table.rows[2].cells
    cells_line4 = table.rows[3].cells
    cells_line5 = table.rows[4].cells
    cells_line6 = table.rows[5].cells
    cells_line7 = table.rows[6].cells
    cells_line1[0].add_paragraph('姓名：').alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line1[1].add_paragraph(patient_name).alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line1[2].add_paragraph('性别：').alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line1[3].add_paragraph(patient_sex).alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line1[4].add_paragraph('年龄').alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line1[5].add_paragraph(patient_age).alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line2[0].add_paragraph('就诊时间').alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line2[2].add_paragraph(diagnosis_time).alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line2[4].add_paragraph('联系方式').alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line2[5].add_paragraph(patient_tel).alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line3[0].add_paragraph('过往病史').alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line3[2].add_paragraph(patient_past_history).alignment = WD_ALIGN_PARAGRAPH.LEFT
    cells_line4[0].add_paragraph('诊断医师').alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line4[2].add_paragraph(doc_name).alignment = WD_ALIGN_PARAGRAPH.LEFT
    cells_line5[0].add_paragraph('症状').alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line5[2].add_paragraph(symptom).alignment = WD_ALIGN_PARAGRAPH.LEFT
    imgPara_w = cells_line6[0].add_paragraph('可疑部位：')
    imgPara_w.alignment = WD_ALIGN_PARAGRAPH.LEFT
    imgPara_i = cells_line6[0].add_paragraph()
    imgPara_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for img_path in img_path_list:
        if img_path is not None and img_path != "":
            imgPara_i.add_run().add_picture(img_path, width=Inches(2.0))
    cells_line7[0].add_paragraph('医师意见：').alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells_line7[2].add_paragraph(diagnosis_content).alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.save(save_name)


def object_to_json(obj):
    # exclude = ["_state"]
    # print(exclude)
    return dict([(kk, obj.__dict__[kk]) for kk in obj.__dict__.keys() if kk != "_state"])


def objects_to_json(objs):
    json_list = []
    for obj in objs:
        json_dict = model_to_dict(obj)
        json_list.append(json_dict)
    return json_list


def list_to_json(objs):
    json_list = []
    for obj in objs:
        json_dict = obj
        json_list.append(json_dict)
    return json_list
#
# @require_http_methods(["POST"])
# def add_book(request):
#     response = {}
#     try:
#         book = Book(book_name=request.GET.get('book_name'))
#         book.save()
#         response['msg'] = 'success'
#         response['error_num'] = 0
#     except Exception as e:
#         response['msg'] = str(e)
#         response['error_num'] = 1
#
#     return JsonResponse(response)
